from mydifflib import ndiff
import docx
from docx.shared import RGBColor
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
import datetime
# file_1='''I !@#$%^1
# am !@#$%^2
# a !@#$%^3
# boy !@#$%^4'''
#
# file_2='''I !@#$%^1
# am not !@#$%^2
# a !@#$%^2
# boy !@#$%^4'''

# src_file="../files/src.txt"
# tar_file="../files/tar.txt"
# src=open(src_file,"r").readlines()
# tar=open(tar_file,"r").readlines()

src_file="../files/指引文件.docx"
tar_file="../files/基金合同样本.docx"
src_data=docx.Document(src_file)
tar_data=docx.Document(tar_file)
src=[p.text + "!@#$%^1" for p in src_data.paragraphs]
tar=[p.text + "!@#$%^1" for p in tar_data.paragraphs]


#src=file_1.split()
#tar=file_2.split()
diff =ndiff(src,tar)
for i in diff:
    print(i)

def get_content_page(text,sep="!@#$%^"):
    temp=text.split(sep)
    content=temp[0][2:]
    page_number=temp[1].replace("\n","")
    return (content,page_number)



def real_change_content( change):
    if "+" in change.keys() and "-" in change.keys():
        if change["+"][0]!=change["-"][0]:
            res={"type":"change"}
            #print("change:")
            temp_diff = ndiff(change["-"][0], change["+"][0])
            tar_mask=""
            src_mask=""
            for c in temp_diff:
                if c.startswith(" "):
                    tar_mask+=" "
                    src_mask+=" "
                    continue
                if c.startswith("-"):
                    src_mask+="-"
                    continue
                if c.startswith("+"):
                    tar_mask+="+"
            temp_diff=ndiff([change["-"][0]],[change["+"][0]])
            # #print("\n".join(temp_diff))
            for ix,ct in enumerate(temp_diff):
                if ct.startswith("-"):
                    res["src_content"]=change["-"][0]
                    res["src_page"] = change["-"][1]
                    continue
                if ct.startswith("+"):
                    res["tar_content"]=change["+"][0]
                    res["tar_page"] = change["+"][1]
                    continue
            res["src_mask"]=src_mask
            res["tar_mask"]=tar_mask

            #     if ct.startswith("?"):
            #         if not "tar_content" in res.keys():
            #             res["src_mask"]=ct[2:].replace("^","-").replace("\n","")
            #         else:
            #             res["tar_mask"]=ct[2:].replace("^","+").replace("\n","")
            # if not "src_mask" in res.keys():
            #     res["src_mask"]=" "* len(res["src_content"])
            # if not "tar_mask" in res.keys():
            #     res["tar_mask"]=" "* len(res["tar_content"])
            #
            # if len(res["src_content"])>len(res["src_mask"]):
            #     res["src_mask"] += " "*(len(res["src_content"])-len(res["src_mask"]))
            # if len(res["tar_content"])>len(res["tar_mask"]):
            #     res["tar_mask"] += " " * (len(res["tar_content"])-len(res["tar_mask"]))
            return res
    elif "+" in change.keys() and not "-" in change.keys():
        #print("add:")
        #print(change["+"][0])
        #print("at page "+change["+"][1])
        res={"type":"add",
             "src_content":"",
             "src_mask":"",
             "src_page":"",
             "tar_content":change["+"][0],
             "tar_mask":"+"*len(change["+"][0]),
             "tar_page":change["+"][1]
             }
        return res
    elif not "+" in change.keys() and "-" in change.keys():
        #print("delete:")
        #print(change["-"][0])
        #print("at page " + change["-"][1])
        res = {"type": "delete",
               "src_content": change["-"][0],
               "src_mask": "-" * len(change["-"][0]),
               "src_page": change["-"][1],
               "tar_content":"",
               "tar_mask": "",
               "tar_page": ""
               }
        return res

def filter_page(diff_res):
    changes=[]
    change={}
    for i in diff_res:
        if i.startswith(" "):
            if len(change) > 0:
                changes.append(real_change_content(change))
                change = {}
        elif i.startswith("?"):
            "do nothing"
        elif i.startswith("+"):
            change["+"]=get_content_page(i)
            changes.append(real_change_content(change))
            change = {}
        elif i.startswith("-"):
            if "-" in change.keys():
                changes.append(real_change_content(change))
                change = {}
                change["-"]=get_content_page(i)
            else:
                change["-"]=get_content_page(i)
    if len(change)>0 :
        changes.append(real_change_content(change))
        change = {}
    final_changes=[i for i in changes if i is not None]
    return final_changes

print("——————————————————————————")
diff =ndiff(src,tar)
changes=filter_page(diff)

for ix,change in enumerate(changes):
    print(ix,change)




def set_cell_v_alignment(cell):
    cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
def set_paragraph_h_alignment(paragraph):
    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER




def render_run_by_charmask(run,charmask):
    if charmask==" ":
        return run
    if charmask=="+":
        run.bold = True
        run.underline=True
        run.font.color.rgb=RGBColor(0x42, 0x24, 0xE9)
    if charmask=="-":
        run.bold=True
        run.font.strike=True
        run.font.color.rgb=RGBColor(0x42, 0x24, 0xE9)



def get_segments(mask):
    segments = []
    for i, mark in enumerate(mask):
        if segments == [] or segments[-1]["mask"] != mark:
            segments.append({"start_pos": i,
                             "stop_pos": i + 1,
                             "mask": mark
                             })
        else:
            segments[-1]["stop_pos"] += 1
    return segments



def add_masked_to_paragraph(text,mask,pg):
    segments = get_segments(mask)
    for segment in segments:
        run = pg.add_run()
        run.text = text[segment["start_pos"]:segment["stop_pos"]]
        render_run_by_charmask(run, segment["mask"])

    #for word,word_mask in zip(text,mask):
    #    #print(word)
    #    run=pg.add_run()
    #    run.text=word
    #    render_run_by_charmask(run,word_mask)

def write_change_to_table(row_no, change ,table):
    table.rows[row_no].cells[0].text=change["src_page"]
    set_cell_v_alignment(table.rows[row_no].cells[0])
    set_paragraph_h_alignment(table.rows[row_no].cells[0].paragraphs[0])
    add_masked_to_paragraph(change["src_content"], change["src_mask"],table.rows[row_no].cells[1].paragraphs[0])
    table.rows[row_no].cells[2].text = change["tar_page"]
    set_cell_v_alignment(table.rows[row_no].cells[2])
    set_paragraph_h_alignment(table.rows[row_no].cells[2].paragraphs[0])
    add_masked_to_paragraph(change["tar_content"], change["tar_mask"], table.rows[row_no].cells[3].paragraphs[0])
    table.rows[row_no].cells[4].text = get_comment(change)

from multiprocessing.dummy import Pool
from itertools import repeat
import re
def get_comment(change):
    if change["type"]=="delete":
        if len(change["src_content"])==0 :
            return "删除空行"
    if change["type"]=="add":
        if len(change["tar_content"])==0:
            return "添加空行"
    return ""


def create_changes_docx(changes):
    document = Document()
    document.sections[0].orientation =docx.enum.section.WD_ORIENT.LANDSCAPE
    new_width, new_height = document.sections[0].page_height, document.sections[0].page_width
    document.sections[0].page_width = new_width
    document.sections[0].page_height = new_height
    r = len(changes)+1  # Number of rows you want
    c = 5  # Number of collumns you want
    table = document.add_table(rows=r, cols=c)
    table.style = 'TableGrid'
    table.columns[0].width = 800000
    table.columns[1].width = 800000
    table.columns[2].width = 800000
    table.columns[3].width = 800000
    table.columns[4].width = 800000
    table.rows[0].cells[0].text="模板页码"
    #table.rows[0].cells[0].width=800000
    table.rows[0].cells[1].text = "模板内容"
    table.rows[0].cells[2].text="修改后页码"
    #table.rows[0].cells[2].width = 800000
    table.rows[0].cells[3].text="修改后内容"
    table.rows[0].cells[4].text = "修改意见"
    #table.rows[0].cells[4].width = 800000
    print("start editting table")
    table_cells = table._cells

    def write_change_to_row(row_no, change):
        row_cells = table_cells[row_no * c:(row_no + 1) * c]
        #c0
        #table.rows[row_no].cells[0].text = change["src_page"]
        row_cells[0].text = change["src_page"]
        #set_cell_v_alignment(table.rows[row_no].cells[0])
        set_cell_v_alignment(row_cells[0])
        #set_paragraph_h_alignment(table.rows[row_no].cells[0].paragraphs[0])
        set_paragraph_h_alignment(row_cells[0].paragraphs[0])
        #c1
        #add_masked_to_paragraph(change["src_content"], change["src_mask"], table.rows[row_no].cells[1].paragraphs[0])
        add_masked_to_paragraph(change["src_content"], change["src_mask"], row_cells[1].paragraphs[0])
        #c2
        #table.rows[row_no].cells[2].text = change["tar_page"]
        row_cells[2].text = change["tar_page"]
        #set_cell_v_alignment(table.rows[row_no].cells[2])
        set_cell_v_alignment(row_cells[2])
        #set_paragraph_h_alignment(table.rows[row_no].cells[2].paragraphs[0])
        set_paragraph_h_alignment(row_cells[2].paragraphs[0])
        #c3
        #add_masked_to_paragraph(change["tar_content"], change["tar_mask"], table.rows[row_no].cells[3].paragraphs[0])
        add_masked_to_paragraph(change["tar_content"], change["tar_mask"], row_cells[3].paragraphs[0])
        #table.rows[row_no].cells[4].text=get_comment(change)
        row_cells[4].text = get_comment(change)
    for row_no, change in enumerate(changes, start=1):
        write_change_to_row(row_no, change)

    #with Pool(32) as pool:
        #pool.starmap(write_change_to_table, zip(range(1, len(changes)+1), changes, repeat(table)))
        #pool.starmap(write_change_to_row, zip(range(1, len(changes)+1), changes))

    return document


print(datetime.datetime.now())
doc=create_changes_docx(changes)
print(datetime.datetime.now())
doc.save('''C:/Users/wyx83/PycharmProjects/word-diff/files/t.docx''')
print(datetime.datetime.now())