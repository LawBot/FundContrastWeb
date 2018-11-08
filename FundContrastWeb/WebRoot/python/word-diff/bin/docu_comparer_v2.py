import os
import sys
import docx
sys.path.append("../reference_code/")
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from mydifflib import ndiff
from docx.shared import RGBColor


def get_content_page(text, sep="!@#$%^"):
    temp = text.split(sep)
    content = temp[0][2:]
    page_number = temp[1].replace("\n", "")
    return (content, page_number)


def real_change_content(change):
    if "+" in change.keys() and "-" in change.keys():
        if change["+"][0] != change["-"][0]:
            res = {"type": "change"}
            temp_diff = ndiff(change["-"][0], change["+"][0])
            tar_mask = ""
            src_mask = ""
            for c in temp_diff:
                if c.startswith(" "):
                    tar_mask += " "
                    src_mask += " "
                    continue
                if c.startswith("-"):
                    src_mask += "-"
                    continue
                if c.startswith("+"):
                    tar_mask += "+"
            temp_diff = ndiff([change["-"][0]], [change["+"][0]])
            for ix, ct in enumerate(temp_diff):
                if ct.startswith("-"):
                    res["src_content"] = change["-"][0]
                    res["src_page"] = change["-"][1]
                    continue
                if ct.startswith("+"):
                    res["tar_content"] = change["+"][0]
                    res["tar_page"] = change["+"][1]
                    continue
            res["src_mask"] = src_mask
            res["tar_mask"] = tar_mask
            return res
    elif "+" in change.keys() and not "-" in change.keys():
        res = {"type": "add",
               "src_content": "",
               "src_mask": "",
               "src_page": "",
               "tar_content": change["+"][0],
               "tar_mask": "+"*len(change["+"][0]),
               "tar_page": change["+"][1]
               }
        return res
    elif not "+" in change.keys() and "-" in change.keys():
        res = {"type": "delete",
               "src_content": change["-"][0],
               "src_mask": "-" * len(change["-"][0]),
               "src_page": change["-"][1],
               "tar_content": "",
               "tar_mask": "",
               "tar_page": ""
               }
        return res


def filter_page(diff_res):
    changes = []
    change = {}
    for i in diff_res:
        if i.startswith(" "):
            if len(change) > 0:
                changes.append(real_change_content(change))
                change = {}
        elif i.startswith("?"):
            "do nothing"
        elif i.startswith("+"):
            change["+"] = get_content_page(i)
            changes.append(real_change_content(change))
            change = {}
        elif i.startswith("-"):
            if "-" in change.keys():
                changes.append(real_change_content(change))
                change = {}
                change["-"] = get_content_page(i)
            else:
                change["-"] = get_content_page(i)
    if len(change) > 0:
        changes.append(real_change_content(change))
        change = {}
    final_changes = [i for i in changes if i is not None]
    return final_changes


def set_cell_v_alignment(cell):
    cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER


def set_row_v_alignment(row):
    for cell in row.cells:
        set_cell_v_alignment(cell)


def set_column_width(column, width):
    for cell in column.cells:
        cell.width = width


def render_run_by_charmask(run, charmask):
    if charmask == " ":
        return run
    if charmask == "+":
        run.bold = True
        run.underline = True
        run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
    if charmask == "-":
        run.bold = True
        run.font.strike = True
        run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)


def get_segments(mask):
    segments = []
    for i, mark in enumerate(mask):
        if segments == [] or segments[-1]["mask"] != mark:
            segments.append({"start_pos": i,
                             "stop_pos": i + 1,
                             "mask": mark
                             })
        else:
            segments[-1]["stop_pos"] += 1
    return segments


def add_masked_to_paragraph(text, mask, pg):
    segments = get_segments(mask)
    for segment in segments:
        run = pg.add_run()
        run.text = text[segment["start_pos"]:segment["stop_pos"]]
        render_run_by_charmask(run, segment["mask"])


def get_comment(change):
    if change["type"] == "delete":
        if len(change["src_content"]) == 0:
            return "删除空行"
    if change["type"] == "add":
        if len(change["tar_content"]) == 0:
            return "添加空行"
    return ""


def set_document_font(document, font):
    document.styles['Normal'].font.name = font
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    return document


def create_changes_docx(changes, source_file, target_file):
    document = Document()
    document.sections[0].orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    new_width, new_height = document.sections[0].page_height, document.sections[0].page_width
    document.sections[0].page_width = new_width
    document.sections[0].page_height = new_height
    # 对比文件信息说明
    source_txt = "修改前文件名：%s  文件大小：%d字节" % (os.path.basename(source_file), os.path.getsize(source_file))
    target_text = "修改后文件名：%s  文件大小：%d字节" % (os.path.basename(target_file), os.path.getsize(target_file))
    document.add_paragraph(source_txt, style='List Bullet')
    document.add_paragraph(target_text, style='List Bullet')
    # 表格内容书写
    r = len(changes)+1  # Number of rows you want
    c = 4  # Number of collumns you want
    table = document.add_table(rows=r, cols=c)
    table.rows[0].cells[0].text = "所在章节"
    table.rows[0].cells[1].text = "修改前内容"
    # table.rows[0].cells[2].text = "修改后页码"
    table.rows[0].cells[2].text = "修改后内容"
    table.rows[0].cells[3].text = "修改意见"
    set_row_v_alignment(table.rows[0])  # 第一行居中显示
    table_cells = table._cells

    def write_change_to_row(row_no, change):
        row_cells = table_cells[row_no * c:(row_no + 1) * c]
        # c0
        row_cells[0].text = change["src_page"] if change["src_page"] else change["tar_page"]
        set_cell_v_alignment(row_cells[0])
        # c1
        add_masked_to_paragraph(change["src_content"], change["src_mask"], row_cells[1].paragraphs[0])
        # c2
        add_masked_to_paragraph(change["tar_content"], change["tar_mask"], row_cells[2].paragraphs[0])
        # c3
        row_cells[3].text = get_comment(change)
    for row_no, change in enumerate(changes, start=1):
        write_change_to_row(row_no, change)
    set_column_width(table.columns[0], new_width*0.16)
    set_column_width(table.columns[1], new_width*0.38)
    set_column_width(table.columns[2], new_width*0.38)
    set_column_width(table.columns[3], new_width*0.1)
    table.style = 'Light Grid Accent 5'
    return set_document_font(document, u"宋体")
