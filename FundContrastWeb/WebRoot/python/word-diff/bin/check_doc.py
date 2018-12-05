import os
import re
import sys
import docx
import time
import pprint
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from word_to_txt import doc2txt
from check_sn import find_sn_exp
from chapter_pattern import add_chapter
from check_content import find_en_punctuation, find_q_str, find_font_exp
from read_config import getConfig

# 获取配置文件
config = getConfig()
punctuation_limit = config.get("punctuation", "punctuation_limit")
db_limit = config.get("double_byte", "db_str")
allow_fonts = [config.get("font_limit", "ch_font").split(";"), config.get("font_limit", "en_font").split(";")]


# 时间戳标记
def timestamp_print(str_):
    time_str = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))
    print("%s %s" % (time_str, str_))


# 发现文档中异常内容
def find_exp(doc_path):
    en_exps = []  # 英文标点异常
    q_exps = []  # 全角字符异常
    sn_exps = []  # 序号异常
    try:
        filename = os.path.split(doc_path)[1]
        if not os.path.exists(doc_path):
            return "-1:文件%s不存在！" % filename
        if os.path.getsize(doc_path) > 50*1024*1024:
            return "-1:文件%s大小超过50M！" % filename
        extension = os.path.splitext(filename)[1]
        if extension.lower() == ".docx" or extension.lower() == ".doc":
            timestamp_print("find font exception")
            font_exps = find_font_exp(doc_path, allow_fonts)  # 字体异常
            txt_path = doc_path.replace(extension, ".txt")
            if doc2txt(doc_path, txt_path):
                content = open(txt_path, "r").readlines()
                content = add_chapter(content)
                timestamp_print("find punctuation&Double-byte exception")
                for line in content:
                    en_exp = find_en_punctuation(line[1], punctuation_limit)
                    q_exp = find_q_str(line[1], db_limit)
                    if en_exp:
                        en_exps.append([line[0], line[1], en_exp])
                    if q_exp:
                        q_exps.append([line[0], line[1], q_exp])
                return [en_exps, q_exps, font_exps]
            else:
                return "-1:文件读取失败！"
        else:
            return "-1:文件非Word文档！"
    except Exception as e:
        return "-1:%s" % e
    finally:
        if "txt_path" in locals() and os.path.exists(txt_path):
            os.remove(txt_path)


# 发现文档异常内容并写入结果文档中
def write_exp_doc(doc_path, exp_path):
    exps = find_exp(doc_path)
    if isinstance(exps, str):
        return exps
    timestamp_print("save result file")
    document = Document()
    document = set_document_font(document, u"宋体")
    document.sections[0].orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    new_width, new_height = document.sections[0].page_height, document.sections[0].page_width
    document.sections[0].page_width = new_width
    document.sections[0].page_height = new_height
    # 1.英文符号异常内容写入
    en_exps = exps[0]
    document.add_paragraph("英文符号异常(数量:%d)" % len(en_exps), style='List Bullet')
    en_table = document.add_table(rows=len(en_exps)+1, cols=2)
    set_column_width(en_table.columns[0], new_width*0.2)
    set_column_width(en_table.columns[1], new_width*0.8)
    write_exp1(en_table, en_exps)
    # 2.全角异常内容写入
    q_exps = exps[1]
    document.add_paragraph("\n")
    document.add_paragraph("全角字符异常(数量:%d)" % len(q_exps), style='List Bullet')
    q_table = document.add_table(rows=len(q_exps)+1, cols=2)
    set_column_width(q_table.columns[0], new_width*0.2)
    set_column_width(q_table.columns[1], new_width*0.8)
    write_exp1(q_table, q_exps)
    # 3.字体异常内容写入
    font_exps = exps[2]
    document.add_paragraph("\n")
    document.add_paragraph("字体异常(数量:%d)" % len(font_exps), style='List Bullet')
    font_table = document.add_table(rows=len(font_exps)+1, cols=2)
    set_column_width(font_table.columns[0], new_width*0.8)
    set_column_width(font_table.columns[1], new_width*0.2)
    write_exp2(font_table, font_exps)
    # 5.保存文档
    document.save(exp_path)
    return "0"


# 写入异常信息表格
def write_exp1(table, exps):
    table.rows[0].cells[0].text = "所在章节"
    table.rows[0].cells[1].text = "异常内容"
    set_row_v_alignment(table.rows[0])  # 第一行居中显示
    table_cells = table._cells
    for row_no, exp in enumerate(exps, 1):
        write_row1(table_cells, row_no, exp)
    table.style = 'Light Grid Accent 5'


# 写入异常信息表格
def write_exp2(table, exps):
    table.rows[0].cells[0].text = "段落内容"
    table.rows[0].cells[1].text = "异常字体"
    set_row_v_alignment(table.rows[0])  # 第一行居中显示
    table_cells = table._cells
    for row_no, exp in enumerate(exps, 1):
        write_row2(table_cells, row_no, exp)
    table.style = 'Light Grid Accent 5'


# 每行异常内容写入
def write_row1(table_cells, row_no, exp):
    row_cells = table_cells[row_no*2:(row_no*2+2)]
    # 第一列
    row_cells[0].text = exp[0]
    set_cell_v_alignment(row_cells[0])
    # 第二列
    para = row_cells[1].paragraphs[0]
    for word in exp[1]:
        word_run = para.add_run()
        word_run.text = word
        # 异常字体加粗标红
        if word in exp[2]:
            word_run.bold = True
            word_run.font.color.rgb = RGBColor(0xFF, 0x0, 0x0)


# 每行异常内容写入
def write_row2(table_cells, row_no, exp):
    row_cells = table_cells[row_no*2:(row_no+1)*2]
    exp_fonts = []
    # 第一列
    para = row_cells[0].paragraphs[0]
    for word in exp:
        word_str = re.sub("\x07|\x06|\x05|\x0c|\r", "", word[0])
        if word_str == "":
            continue
        word_run = para.add_run()
        # pprint.pprint(word_str)
        word_run.text = word_str
        # 异常字体加粗标红
        if word[2]:
            word_run.bold = True
            word_run.font.color.rgb = RGBColor(0xFF, 0x0, 0x0)
            if word[1] not in exp_fonts:
                exp_fonts.append(word[1])
    # 第二列
    row_cells[1].text = ",".join([f for f in exp_fonts if f != ""])


# 设置字体
def set_document_font(document, font):
    document.styles['Normal'].font.name = font
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    return document


# 单元格设置文字居中
def set_cell_v_alignment(cell):
    cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER


# 表格行设置居中
def set_row_v_alignment(row):
    for cell in row.cells:
        set_cell_v_alignment(cell)


# 设置列宽
def set_column_width(column, width):
    for cell in column.cells:
        cell.width = width


if __name__ == '__main__':
    # doc_path = r"C:\Users\huangzhouzhou\Documents\GitHub\word-diff\tmp\e21.doc"
    # result_doc = r"C:\Users\huangzhouzhou\Documents\GitHub\word-diff\tmp\e21_exceptions.docx"
    # print(write_exp_doc(doc_path, result_doc))
    if len(sys.argv) == 3:
        doc_path = sys.argv[1]
        result_doc = sys.argv[2]
        print(write_exp_doc(doc_path, result_doc))
    else:
        print("-1:参数数量异常！")
