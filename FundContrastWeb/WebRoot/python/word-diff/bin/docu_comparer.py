import docx
import sys
sys.path.append("../reference_code/")
from mydifflib import ndiff
from docx.shared import RGBColor
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_COLOR_INDEX


def get_content_page(text, sep="!@#$%^"):
    temp = text.split(sep)
    content = temp[0][2:]
    page_number = temp[1].replace("\n", "")
    return (content, page_number)


def real_change_content(change):
    if "+" in change.keys() and "-" in change.keys():
        if change["+"][0] != change["-"][0]:
            res = {"type": "change"}
            temp_diff = ndiff(change["-"][0], change["+"][0])
            tar_mask = ""
            src_mask = ""
            for c in temp_diff:
                if c.startswith(" "):
                    tar_mask += " "
                    src_mask += " "
                    continue
                if c.startswith("-"):
                    src_mask += "-"
                    continue
                if c.startswith("+"):
                    tar_mask += "+"
            temp_diff = ndiff([change["-"][0]], [change["+"][0]])
            for ix, ct in enumerate(temp_diff):
                if ct.startswith("-"):
                    res["src_content"] = change["-"][0]
                    res["src_page"] = change["-"][1]
                    continue
                if ct.startswith("+"):
                    res["tar_content"] = change["+"][0]
                    res["tar_page"] = change["+"][1]
                    continue
            res["src_mask"] = src_mask
            res["tar_mask"] = tar_mask
            return res
    elif "+" in change.keys() and not "-" in change.keys():
        res = {"type": "add",
               "src_content": "",
               "src_mask": "",
               "src_page": "",
               "tar_content": change["+"][0],
               "tar_mask": "+"*len(change["+"][0]),
               "tar_page": change["+"][1]
               }
        return res
    elif not "+" in change.keys() and "-" in change.keys():
        res = {"type": "delete",
               "src_content": change["-"][0],
               "src_mask": "-" * len(change["-"][0]),
               "src_page": change["-"][1],
               "tar_content": "",
               "tar_mask": "",
               "tar_page": ""
               }
        return res


def filter_page(diff_res):
    changes = []
    change = {}
    for i in diff_res:
        if i.startswith(" "):
            if len(change) > 0:
                changes.append(real_change_content(change))
                change = {}
        elif i.startswith("?"):
            "do nothing"
        elif i.startswith("+"):
            change["+"] = get_content_page(i)
            changes.append(real_change_content(change))
            change = {}
        elif i.startswith("-"):
            if "-" in change.keys():
                changes.append(real_change_content(change))
                change = {}
                change["-"] = get_content_page(i)
            else:
                change["-"] = get_content_page(i)
    if len(change) > 0:
        changes.append(real_change_content(change))
        change = {}
    final_changes = [i for i in changes if i is not None]
    return final_changes


def set_cell_v_alignment(cell):
    cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_h_alignment(paragraph):
    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER


def render_run_by_charmask(run, charmask):
    if charmask == " ":
        return run
    if charmask == "+":
        run.bold = True
        run.underline = True
        run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
    if charmask == "-":
        run.bold = True
        run.font.strike = True
        run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)


def get_segments(mask):
    segments = []
    for i, mark in enumerate(mask):
        if segments == [] or segments[-1]["mask"] != mark:
            segments.append({"start_pos": i,
                             "stop_pos": i + 1,
                             "mask": mark
                             })
        else:
            segments[-1]["stop_pos"] += 1
    return segments


def add_masked_to_paragraph(text, mask, pg):
    segments = get_segments(mask)
    for segment in segments:
        run = pg.add_run()
        run.text = text[segment["start_pos"]:segment["stop_pos"]]
        render_run_by_charmask(run, segment["mask"])


def get_comment(change):
    if change["type"] == "delete":
        if len(change["src_content"]) == 0:
            return "删除空行"
    if change["type"] == "add":
        if len(change["tar_content"]) == 0:
            return "添加空行"
    return ""


def set_section_landspace(section):
    section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    section.page_height, section.page_width = section.page_width, section.page_height


def create_changes_docx(changes):
    document = Document()
    document.sections[0].orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    new_width, new_height = document.sections[0].page_height, document.sections[0].page_width
    document.sections[0].page_width = new_width
    document.sections[0].page_height = new_height
    r = len(changes)+1  # Number of rows you want
    c = 5  # Number of collumns you want
    table = document.add_table(rows=r, cols=c)
    table.style = 'TableGrid'
    table.rows[0].cells[0].text = "模板页码"
    table.rows[0].cells[1].text = "模板内容"
    table.rows[0].cells[2].text = "修改后页码"
    table.rows[0].cells[3].text = "修改后内容"
    table.rows[0].cells[4].text = "修改意见"
    table_cells = table._cells

    def write_change_to_row(row_no, change):
        row_cells = table_cells[row_no * c:(row_no + 1) * c]
        # c0
        row_cells[0].text = change["src_page"]
        set_cell_v_alignment(row_cells[0])
        set_paragraph_h_alignment(row_cells[0].paragraphs[0])
        # c1
        add_masked_to_paragraph(change["src_content"], change["src_mask"], row_cells[1].paragraphs[0])
        # c2
        row_cells[2].text = change["tar_page"]
        set_cell_v_alignment(row_cells[2])
        set_paragraph_h_alignment(row_cells[2].paragraphs[0])
        # c3
        add_masked_to_paragraph(change["tar_content"], change["tar_mask"], row_cells[3].paragraphs[0])
        # c4
        row_cells[4].text = get_comment(change)
    for row_no, change in enumerate(changes, start=1):
        write_change_to_row(row_no, change)

    return document
